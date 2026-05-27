from io import BytesIO
from pathlib import PurePath

from app.services.mime_types import is_excel_document, is_word_document, resolve_mime_type


class DocumentExtractionError(Exception):
    pass


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("python-docx is not installed") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError(f"Failed to read Word document: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))

    return "\n".join(lines)


def extract_xlsx_text(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentExtractionError("openpyxl is not installed") from exc

    try:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise DocumentExtractionError(f"Failed to read Excel workbook: {exc}") from exc

    lines: list[str] = []
    for worksheet in workbook.worksheets:
        lines.append(f"--- Sheet: {worksheet.title} ---")
        for row in worksheet.iter_rows(values_only=True):
            cells = [
                str(value).strip()
                for value in row
                if value is not None and str(value).strip()
            ]
            if cells:
                lines.append("\t".join(cells))

    workbook.close()
    return "\n".join(lines)


def extract_office_document_text(
    *,
    content: bytes,
    mime_type: str,
    filename: str | None = None,
) -> str:
    resolved_mime_type = resolve_mime_type(mime_type, filename)
    extension = PurePath(filename).suffix.lower() if filename else ""

    if extension == ".doc":
        raise DocumentExtractionError(
            "Legacy .doc format is not supported; upload .docx instead",
        )
    if extension == ".xls":
        raise DocumentExtractionError(
            "Legacy .xls format is not supported; upload .xlsx instead",
        )

    if is_word_document(resolved_mime_type, filename):
        text = extract_docx_text(content)
        if not text.strip():
            raise DocumentExtractionError("Word document contains no extractable text")
        return text

    if is_excel_document(resolved_mime_type, filename):
        text = extract_xlsx_text(content)
        if not text.strip():
            raise DocumentExtractionError("Excel workbook contains no extractable text")
        return text

    raise DocumentExtractionError(
        f"Unsupported office document MIME type: {resolved_mime_type}",
    )
