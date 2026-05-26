from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook

from app.services.document_extractor import DocumentExtractionError, extract_office_document_text
from app.services.mime_types import resolve_mime_type
from app.services.ocr_provider import process_document_content


def _build_docx(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_xlsx(rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sheet1"
    for row in rows:
        worksheet.append(row)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def test_resolve_mime_type_from_extension():
    assert (
        resolve_mime_type("application/octet-stream", "report.docx")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert (
        resolve_mime_type("application/octet-stream", "data.xlsx")
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


def test_extract_docx_text_includes_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("Invoice ACME")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Total"
    table.rows[0].cells[1].text = "100"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_office_document_text(
        content=buffer.getvalue(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="invoice.docx",
    )

    assert "Invoice ACME" in text
    assert "Total" in text
    assert "100" in text


def test_extract_xlsx_text_includes_sheet_and_rows():
    content = _build_xlsx([["Counterparty", "Amount"], ["ACME", "250"]])

    text = extract_office_document_text(
        content=content,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="report.xlsx",
    )

    assert "--- Sheet: Sheet1 ---" in text
    assert "Counterparty" in text
    assert "ACME" in text
    assert "250" in text


def test_process_document_content_routes_office_files_to_extractor():
    content = _build_docx("Contract for Beta LLC")

    result = process_document_content(
        content=content,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="contract.docx",
        ocr_provider_name="stub",
    )

    assert result.provider == "office_extractor"
    assert "Contract for Beta LLC" in result.full_text
    assert result.confidence is not None


def test_process_document_content_routes_pdf_to_stub_provider():
    content = b"plain text invoice"

    result = process_document_content(
        content=content,
        mime_type="application/pdf",
        filename="invoice.pdf",
        ocr_provider_name="stub",
    )

    assert result.provider == "stub"
    assert "plain text invoice" in result.full_text


def test_extract_office_document_text_rejects_empty_docx():
    content = _build_docx("")

    with pytest.raises(DocumentExtractionError, match="no extractable text"):
        extract_office_document_text(
            content=content,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="empty.docx",
        )
